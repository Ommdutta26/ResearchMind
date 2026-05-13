"""
utils/export.py
Rich export helpers: PDF (ReportLab) and DOCX (python-docx).
"""

from __future__ import annotations

import os
import re
import time

from core.config import EXPORT_DIR

os.makedirs(EXPORT_DIR, exist_ok=True)


# ── PDF export ────────────────────────────────────────────────────────────────

def export_pdf(summary: str, topic: str) -> str:
    """Generate a formatted PDF report. Returns the file path."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle,
        )

        filename = os.path.join(
            EXPORT_DIR, f"report_{time.strftime('%Y%m%d_%H%M%S')}.pdf"
        )
        doc    = SimpleDocTemplate(filename, pagesize=A4,
                                   leftMargin=2*cm, rightMargin=2*cm,
                                   topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "NexusTitle",
            parent=styles["Title"],
            fontSize=22,
            textColor=colors.HexColor("#1a1a2e"),
            spaceAfter=6,
        )
        heading_style = ParagraphStyle(
            "NexusH2",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=colors.HexColor("#16213e"),
            spaceBefore=12,
            spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "NexusBody",
            parent=styles["BodyText"],
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#2d2d2d"),
        )
        meta_style = ParagraphStyle(
            "NexusMeta",
            parent=styles["Italic"],
            fontSize=9,
            textColor=colors.grey,
        )

        story = []

        # Header block
        story.append(Paragraph("NexusResearch", title_style))
        story.append(Paragraph(f"Topic: {topic}", heading_style))
        story.append(Paragraph(
            f"Generated: {time.strftime('%B %d, %Y at %H:%M')}",
            meta_style,
        ))
        story.append(HRFlowable(width="100%", thickness=1,
                                 color=colors.HexColor("#e0e0e0"), spaceAfter=12))

        # Parse markdown-ish content into paragraphs
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith("#### "):
                story.append(Paragraph(line[5:], heading_style))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:], heading_style))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], heading_style))
            elif line.startswith("- "):
                story.append(Paragraph(f"• {line[2:]}", body_style))
            else:
                # Strip basic markdown bold/italic
                clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
                clean = re.sub(r"\*(.+?)\*",     r"<i>\1</i>", clean)
                story.append(Paragraph(clean, body_style))

        doc.build(story)
        return filename

    except ImportError:
        raise RuntimeError(
            "reportlab is not installed. Run: pip install reportlab"
        )


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_docx(summary: str, topic: str) -> str:
    """Generate a formatted Word document. Returns the file path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import copy

        filename = os.path.join(
            EXPORT_DIR, f"report_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        )
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # Title
        title = doc.add_heading("NexusResearch", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.runs[0]
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Subtitle
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = sub.add_run(f"Topic: {topic}")
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

        meta = doc.add_paragraph(f"Generated: {time.strftime('%B %d, %Y at %H:%M')}")
        meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        meta.runs[0].font.size = Pt(9)

        doc.add_paragraph()  # spacer

        # Body content
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("## "):
                h = doc.add_heading(line[3:], level=2)
                h.runs[0].font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
            elif line.startswith("### "):
                h = doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(line[2:])
            else:
                # Handle **bold** inline
                p  = doc.add_paragraph()
                parts = re.split(r"(\*\*.+?\*\*)", line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        r = p.add_run(part[2:-2])
                        r.bold = True
                    else:
                        p.add_run(part)

        doc.save(filename)
        return filename

    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
